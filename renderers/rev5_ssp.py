"""
Rev 5 SSP renderer.

Produces a Word document with control-implementation statements in the
format FedRAMP Rev 5 SSPs use. Aggregator-backed capabilities contribute
live-state determinations (timestamped, with real numbers); declared-mode
capabilities contribute their hand-written narrative.

Multiple capabilities can roll up under one control — they concatenate.

Run:
    python -m renderers.rev5_ssp --out samples/rev5_ssp_fragment.docx
    python -m renderers.rev5_ssp --out samples/rev5_ssp_fragment.docx --fixtures tests/fixtures
"""

from __future__ import annotations

import argparse
import re
import uuid
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor

from aggregators._base import AggregatorRunContext
from renderers.shared.capability_loader import (
    index_by_rev5_control,
    load_all,
)
from renderers.shared.determinations import DeterminationResolver, ResolvedCapabilityEntry

CONTROL_FAMILIES = {
    "AC": "Access Control",
    "AT": "Awareness and Training",
    "AU": "Audit and Accountability",
    "CA": "Assessment, Authorization, and Monitoring",
    "CM": "Configuration Management",
    "CP": "Contingency Planning",
    "IA": "Identification and Authentication",
    "IR": "Incident Response",
    "MA": "Maintenance",
    "MP": "Media Protection",
    "PE": "Physical and Environmental Protection",
    "PL": "Planning",
    "PM": "Program Management",
    "PS": "Personnel Security",
    "PT": "Personally Identifiable Information Processing and Transparency",
    "RA": "Risk Assessment",
    "SA": "System and Services Acquisition",
    "SC": "System and Communications Protection",
    "SI": "System and Information Integrity",
    "SR": "Supply Chain Risk Management",
}


def _control_family(control_id: str) -> str:
    prefix = re.match(r"^([A-Z]+)", control_id).group(1)
    return CONTROL_FAMILIES.get(prefix, "Unknown")


def _sort_key(control_id: str):
    m = re.match(r"^([A-Z]+)-(\d+)(?:\((\d+)\))?", control_id)
    if not m:
        return (control_id,)
    family, num, enh = m.groups()
    return (family, int(num), int(enh) if enh else 0)


def _rollup_status(entries: list[ResolvedCapabilityEntry]) -> str:
    statuses: set[str] = set()
    for e in entries:
        if e.determination:
            statuses.add(e.determination.status)
        else:
            for ctrl_entry in e.capability.rev5_controls():
                s = ctrl_entry.get("implementation_status")
                if s:
                    statuses.add(s)
    if "Planned" in statuses:
        return "Planned"
    if "Partially Implemented" in statuses:
        return "Partially Implemented"
    if "Inconclusive" in statuses:
        return "Inconclusive"
    if "Alternative Implementation" in statuses:
        return "Alternative Implementation"
    return "Implemented" if statuses else "Not Documented"


def render(
    output_path: Path,
    requested_controls: list[str] | None = None,
    fixtures_dir: str | None = None,
    strict_freshness: bool = False,
) -> Path:
    caps = load_all()
    idx = index_by_rev5_control(caps)

    if requested_controls:
        idx = {c: lst for c, lst in idx.items() if c in requested_controls}

    ctx = AggregatorRunContext(
        fixture_mode=bool(fixtures_dir),
        fixture_dir=fixtures_dir or "",
        strict_freshness=strict_freshness,
        run_id=str(uuid.uuid4()),
    )
    resolver = DeterminationResolver(ctx)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    title = doc.add_heading("System Security Plan — Control Implementation Statements", 0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    p = doc.add_paragraph()
    p.add_run(
        "Generated from canonical capability definitions. Aggregator-backed "
        "capabilities reflect observed live state at render time; declared "
        "capabilities reflect the engineering team's authored statement. Do "
        "not edit this document directly — modify the source and regenerate."
    ).italic = True

    by_family: dict[str, list[str]] = {}
    for ctrl in sorted(idx.keys(), key=_sort_key):
        by_family.setdefault(_control_family(ctrl), []).append(ctrl)

    for family in sorted(by_family.keys()):
        doc.add_heading(family, level=1)
        for control_id in by_family[family]:
            entries = resolver.for_rev5_control(control_id, idx[control_id])
            _render_control(doc, control_id, entries)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def _render_control(doc, control_id: str, entries: list[ResolvedCapabilityEntry]):
    doc.add_heading(control_id, level=2)

    status = _rollup_status(entries)
    contributing = sorted({e.capability.id for e in entries})

    table = doc.add_table(rows=2, cols=2)
    table.style = "Light Grid Accent 1"
    table.cell(0, 0).text = "Implementation Status"
    table.cell(0, 1).text = status
    table.cell(1, 0).text = "Contributing Capabilities"
    table.cell(1, 1).text = ", ".join(contributing)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                if para.runs:
                    para.runs[0].font.size = Pt(10)

    for entry in entries:
        cap = entry.capability
        flavor = "live-state" if entry.determination else "declared"
        heading = doc.add_heading(f"{cap.id}", level=3)
        for r in heading.runs:
            r.font.size = Pt(11)
        p = doc.add_paragraph()
        p.add_run(f"({flavor}) ").italic = True

        if entry.determination:
            doc.add_paragraph(entry.determination.statement)
            ts = doc.add_paragraph()
            ts.add_run(f"Observed at: {entry.determination.observed_at}").italic = True
        else:
            doc.add_paragraph(entry.declared_statement)

    # Evidence section
    doc.add_heading("Validation Evidence", level=3)
    for entry in entries:
        if entry.determination:
            for ref in entry.determination.evidence_refs:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{ref}").bold = True
                p.add_run(f" — gathered by {entry.capability.aggregator_path}")
        else:
            for ev in entry.capability.evidence():
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{ev['id']}").bold = True
                p.add_run(
                    f" — {ev['type']} from {ev['source']}, "
                    f"{ev.get('schedule', 'on-demand')}, "
                    f"{ev.get('validation_method', 'automated')}"
                )

    # Provenance
    doc.add_heading("Provenance", level=3)
    for entry in entries:
        prov = entry.capability.provenance()
        bits = [
            f"FRMR: {prov.get('frmr_version', 'n/a')}",
            f"last reviewed {prov.get('last_reviewed', 'n/a')}",
        ]
        for a in prov.get("validated_in_assessment", []):
            bits.append(f"validated in {a['csp']}/{a['assessor_3pao']} ({a['date']})")
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{entry.capability.id}: ").bold = True
        p.add_run("; ".join(bits))

    doc.add_paragraph()


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", required=True, type=Path)
    ap.add_argument("--controls", default=None,
                    help="Comma-separated list of controls; default = all that have capabilities")
    ap.add_argument("--fixtures", default=None,
                    help="Path to a fixtures directory — runs aggregators in fixture mode")
    ap.add_argument("--strict-freshness", action="store_true")
    args = ap.parse_args()

    controls = [c.strip() for c in args.controls.split(",")] if args.controls else None
    out = render(args.out, controls, fixtures_dir=args.fixtures, strict_freshness=args.strict_freshness)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
